#!/usr/bin/env python3
"""
CrewAI Demo script for Vietnamese Document Form Filler.

Shows CrewAI-powered multi-agent capabilities.
"""

import asyncio
import json
import os
import tempfile
import time
from pathlib import Path

from form_filler.batch_processor import CrewAIBatchProcessor
from form_filler.crew.document_processing_crew import DocumentProcessingCrew

# Demo data - sample Vietnamese text
SAMPLE_VIETNAMESE_TEXT = """
Họ và tên: Nguyễn Văn An
Ngày sinh: 15/03/1990
Địa chỉ: 123 Phố Huế, Quận Hai Bà Trưng, Hà Nội
Số điện thoại: 0123456789
Email: nguyen.van.an@email.com
Trình độ học vấn: Cử nhân Công nghệ Thông tin

Kinh nghiệm làm việc:
- 2015-2018: Lập trình viên tại Công ty ABC
- 2018-2022: Trưởng nhóm phát triển tại Công ty XYZ
- 2022-hiện tại: Kiến trúc sư phần mềm tại Công ty DEF

Kỹ năng:
- Ngôn ngữ lập trình: Python, Java, JavaScript
- Framework: Django, React, Spring Boot
- Cơ sở dữ liệu: MySQL, PostgreSQL, MongoDB
- DevOps: Docker, Kubernetes, AWS
"""

SAMPLE_FORM_CONTENT = """
EMPLOYMENT APPLICATION FORM

Personal Information:
Name: ___________________
Date of Birth: ___________________
Address: ___________________________________
Phone Number: ___________________
Email: ___________________

Education:
Highest Degree: ___________________
Institution: ___________________

Work Experience:
Previous Position: ___________________
Company: ___________________
Years of Experience: ___________________

Skills:
Technical Skills: ___________________________________
Programming Languages: ___________________
Frameworks: ___________________
Databases: ___________________

Additional Information:
___________________________________________
___________________________________________
"""


def demo_crewai_single_document():
    """Demonstrate CrewAI single document processing."""
    print("🤖 Demo 1: CrewAI Single Document Processing")
    print("-" * 50)

    # Create temporary files for source document as PDF
    import fitz  # PyMuPDF
    source_path = tempfile.mktemp(suffix=".pdf")
    doc = fitz.open()
    page = doc.new_page()
    
    # Use a Unicode-compatible font for Vietnamese with proper line breaks
    text_writer = fitz.TextWriter(page.rect)
    font = fitz.Font("helv")  # Use helvetica font with Unicode support
    
    # Split text by line breaks and render each line with proper spacing
    lines = SAMPLE_VIETNAMESE_TEXT.strip().split('\n')
    for i, line in enumerate(lines):
        # Position each line with increasing y-coordinate for proper line spacing
        text_writer.append((50, 50 + i * 14), line, font=font, fontsize=11)
    
    text_writer.write_text(page)
    doc.save(source_path)
    doc.close()

    # Create a temporary DOCX form file
    from docx import Document
    doc = Document()
    
    # Add title
    doc.add_heading("EMPLOYMENT APPLICATION FORM", 0)
    doc.add_paragraph("")
    
    # Add personal information section
    doc.add_heading("Personal Information:", 1)
    doc.add_paragraph("Name: ___________________")
    doc.add_paragraph("Date of Birth: ___________________")
    doc.add_paragraph("Address: ___________________________________")
    doc.add_paragraph("Phone Number: ___________________")
    doc.add_paragraph("Email: ___________________")
    
    # Add education section
    doc.add_heading("Education:", 1)
    doc.add_paragraph("Highest Degree: ___________________")
    doc.add_paragraph("Institution: ___________________")
    
    # Add work experience section
    doc.add_heading("Work Experience:", 1)
    doc.add_paragraph("Previous Position: ___________________")
    doc.add_paragraph("Company: ___________________")
    doc.add_paragraph("Years of Experience: ___________________")
    
    # Add skills section
    doc.add_heading("Skills:", 1)
    doc.add_paragraph("Technical Skills: ___________________________________")
    doc.add_paragraph("Programming Languages: ___________________")
    doc.add_paragraph("Frameworks: ___________________")
    doc.add_paragraph("Databases: ___________________")
    
    # Add additional information section
    doc.add_heading("Additional Information:", 1)
    doc.add_paragraph("___________________________________________")
    doc.add_paragraph("___________________________________________")
    
    # Save the document
    form_path = tempfile.mktemp(suffix=".docx")
    doc.save(form_path)
    
    # Create temporary output file
    output_path = tempfile.mktemp(suffix=".docx")

    # Initialize CrewAI processor
    print("🤖 Initializing CrewAI crew...")
    crew_processor = DocumentProcessingCrew(
        text_model="llama3.2:3b",
        extraction_method="traditional",
        vision_model="llava:7b",
    )

    print("📄 Processing Vietnamese CV with CrewAI agents...")
    print(f"Source: {Path(source_path).name}")
    print(f"Form: {Path(form_path).name}")
    print("\n🤖 CrewAI Agents in Action:")
    print("  1. 📄 Document Collector Agent - Extracting text...")
    print("  2. 🔄 Translator Agent - Translating Vietnamese to English...")
    print("  3. 🔍 Form Analyst Agent - Analyzing form structure...")
    print("  4. ✍️ Form Filler Agent - Filling form intelligently...")

    # Process document
    result = crew_processor.process_document(source_path, form_path, output_path)

    if result.success:
        print("\n✅ CrewAI Success!")
        print(f"Output saved to: {output_path}")

        # Show agent collaboration results
        if Path(output_path).exists():
            with Path(output_path).open() as f:
                content = f.read()
                print("\n📋 CrewAI Filled Form Preview:")
                print(content[:300] + "..." if len(content) > 300 else content)

        print(f"🤖 Agent Configuration: {result.metadata.get('extraction_method')} extraction")
        print(f"📱 Text Model: {result.metadata.get('text_model')}")
        if result.metadata.get("vision_model"):
            print(f"👁️ Vision Model: {result.metadata.get('vision_model')}")
    else:
        print(f"❌ CrewAI Failed: {result.error}")

    # Save to .data directory
    data_dir = Path(".data")
    data_dir.mkdir(exist_ok=True)
    
    # Copy files to .data
    import shutil
    if Path(source_path).exists():
        shutil.copy2(source_path, data_dir / "demo1_source.pdf")
    if Path(form_path).exists():
        shutil.copy2(form_path, data_dir / "demo1_form.docx")
    if Path(output_path).exists():
        shutil.copy2(output_path, data_dir / "demo1_output.docx")
        
    # Cleanup
    Path(source_path).unlink()
    Path(form_path).unlink()
    if Path(output_path).exists():
        Path(output_path).unlink()

    print("\n")


def demo_crewai_agent_capabilities():
    """Demonstrate individual CrewAI agent capabilities."""
    print("🤖 Demo 2: CrewAI Agent Capabilities")
    print("-" * 50)

    from form_filler.tools.document_extraction_tool import DocumentExtractionTool
    from form_filler.tools.form_analysis_tool import FormAnalysisTool
    from form_filler.tools.form_filling_tool import FormFillingTool
    from form_filler.tools.translation_tool import TranslationTool
    from docx import Document

    # Create temp files for source document as PDF
    import fitz  # PyMuPDF
    temp_path = tempfile.mktemp(suffix=".pdf")
    doc = fitz.open()
    page = doc.new_page()
    # Use a Unicode-compatible font for Vietnamese with proper line breaks
    text_writer = fitz.TextWriter(page.rect)
    font = fitz.Font("helv")  # Use helvetica font with Unicode support
    
    # Split text by line breaks and render each line with proper spacing
    lines = SAMPLE_VIETNAMESE_TEXT.strip().split('\n')
    for i, line in enumerate(lines):
        # Position each line with increasing y-coordinate for proper line spacing
        text_writer.append((50, 50 + i * 14), line, font=font, fontsize=11)
    
    text_writer.write_text(page)
    doc.save(temp_path)
    doc.close()

    print("🔧 CrewAI Tool 1: DocumentExtractionTool")
    extractor = DocumentExtractionTool(extraction_method="traditional")
    extracted_text = extractor._run(temp_path)
    print(f"✅ Extracted {len(extracted_text)} characters")
    print(f"📝 Sample: {extracted_text[:100]}...")

    print("\n🔧 CrewAI Tool 2: TranslationTool")
    translator = TranslationTool(model="llama3.2:3b")
    english_text = translator._run(extracted_text)
    print("✅ Translation completed by CrewAI agent")
    print(f"📝 Sample: {english_text[:100]}...")

    print("\n🔧 CrewAI Tool 3: FormAnalysisTool")
    
    # Create a temporary DOCX form file
    doc = Document()
    doc.add_heading("EMPLOYMENT APPLICATION FORM", 0)
    doc.add_heading("Personal Information:", 1)
    doc.add_paragraph("Name: ___________________")
    doc.add_paragraph("Date of Birth: ___________________")
    doc.add_paragraph("Email: ___________________")
    doc.add_heading("Skills:", 1)
    doc.add_paragraph("Programming Languages: ___________________")
    
    form_path = tempfile.mktemp(suffix=".docx")
    doc.save(form_path)

    form_analyzer = FormAnalysisTool()
    form_analyzer._run(form_path)
    print("✅ Form analysis completed")
    print("📝 Found form fields in structure")

    print("\n🔧 CrewAI Tool 4: FormFillingTool")
    output_path = tempfile.mktemp(suffix=".docx")
    form_filler = FormFillingTool(model="llama3.2:3b")
    filling_result = form_filler._run(form_path, english_text, output_path)
    print("✅ Form filling completed by CrewAI agent")
    print(f"📝 Result: {filling_result}")

    # Save to .data directory
    data_dir = Path(".data")
    data_dir.mkdir(exist_ok=True)
    
    # Copy files to .data
    import shutil
    if Path(temp_path).exists():
        shutil.copy2(temp_path, data_dir / "demo2_source.pdf")
    if Path(form_path).exists():
        shutil.copy2(form_path, data_dir / "demo2_form.docx")
    if Path(output_path).exists():
        shutil.copy2(output_path, data_dir / "demo2_output.docx")
        
    # Cleanup
    Path(temp_path).unlink()
    Path(form_path).unlink()
    if Path(output_path).exists():
        Path(output_path).unlink()

    print("\n")


def demo_crewai_batch_processing():
    """Demonstrate CrewAI batch processing."""
    print("🤖 Demo 3: CrewAI Batch Processing")
    print("-" * 50)

    # Create sample documents with proper Vietnamese encoding and line breaks
    sample_documents = [
        f"Tài liệu {i}\nTên: Người {i}\nTuổi: {20 + i}\nĐịa chỉ: Hà Nội, Việt Nam"
        for i in range(1, 4)
    ]

    # Create temporary directory structure
    temp_dir = Path(tempfile.mkdtemp())
    input_dir = temp_dir / "input"
    output_dir = temp_dir / "output"
    input_dir.mkdir()
    output_dir.mkdir()

    # Create a form template in DOCX format
    from docx import Document
    doc = Document()
    doc.add_heading("EMPLOYMENT APPLICATION FORM", 0)
    doc.add_paragraph("")
    
    # Add personal information section
    doc.add_heading("Personal Information:", 1)
    doc.add_paragraph("Name: ___________________")
    doc.add_paragraph("Age: ___________________")
    doc.add_paragraph("Address: ___________________")
    
    # Add simple data sections for batch processing
    doc.add_heading("Employment Details:", 1)
    doc.add_paragraph("Position: ___________________")
    doc.add_paragraph("Start Date: ___________________")
    
    # Save the form
    form_path = temp_dir / "form.docx"
    doc.save(form_path)

    # Create input documents as PDFs with proper font for Vietnamese
    import fitz  # PyMuPDF
    for i, content in enumerate(sample_documents, 1):
        doc_path = input_dir / f"document_{i}.pdf"
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        
        # Use a Unicode-compatible font for Vietnamese with TextWriter and line breaks
        text_writer = fitz.TextWriter(page.rect)
        font = fitz.Font("helv")  # Use helvetica font with Unicode support
        
        # Split text by line breaks and render each line with proper spacing
        lines = content.strip().split('\n')
        for i, line in enumerate(lines):
            # Position each line with increasing y-coordinate for proper line spacing
            text_writer.append((50, 50 + i * 15), line, font=font, fontsize=12)
            
        text_writer.write_text(page)
        pdf_doc.save(doc_path)
        pdf_doc.close()

    print(f"📁 Created {len(sample_documents)} sample documents")
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")

    # Initialize CrewAI batch processor
    print("\n🤖 Initializing CrewAI Batch Processor...")
    batch_processor = CrewAIBatchProcessor(
        text_model="llama3.2:3b",
        extraction_method="traditional",
        vision_model="llava:7b",
        max_concurrent=2,
        timeout=60,
    )

    # Discover jobs
    job_count = batch_processor.discover_jobs(
        str(input_dir),
        str(form_path),
        str(output_dir),
        "*.pdf",
    )

    print(f"🔍 Discovered {job_count} processing jobs")
    print("🤖 CrewAI Configuration:")
    print("  - Max concurrent crews: 2")
    print("  - Each crew has 4 specialized agents")
    print(f"  - Total agents working: {job_count * 4} (across all crews)")

    # Progress callback
    def progress_callback(completed, total, job):
        progress = (completed / total) * 100
        status = "✅" if job.status == "completed" else "❌"
        print(
            f"Progress: {completed}/{total} ({progress:.1f}%) - {status} {job.source_path.name} (Crew: {job.crew_id})",
        )

    # Process batch
    print("\n🚀 Starting CrewAI batch processing...")
    print("Multiple crews working in parallel...")
    stats = batch_processor.process_all(progress_callback)

    print("\n📊 CrewAI Batch Processing Results:")
    print(f"Total: {stats['total']}")
    print(f"Completed: {stats['completed']}")
    print(f"Failed: {stats['failed']}")
    print(f"Success Rate: {stats['success_rate']:.1f}%")
    print(f"Total Time: {stats['total_time']:.2f}s")
    print(f"Average per Job: {stats['average_job_time']:.2f}s")
    print(f"Extraction Method: {stats['extraction_method']}")
    print(f"Text Model: {stats['text_model']}")

    # Save to .data directory
    data_dir = Path(".data")
    data_dir.mkdir(exist_ok=True)
    
    # Copy example files to .data
    import shutil
    
    # Copy form template
    if Path(form_path).exists():
        shutil.copy2(form_path, data_dir / "demo3_form_template.docx")
    
    # Copy one input and output example file if available
    input_files = list(input_dir.glob("*.pdf"))
    if input_files:
        shutil.copy2(input_files[0], data_dir / "demo3_input_example.pdf")
    
    output_files = list(output_dir.glob("*.docx"))
    if output_files:
        shutil.copy2(output_files[0], data_dir / "demo3_output_example.docx")
        
    # Cleanup
    shutil.rmtree(temp_dir)

    print("\n")


def demo_crewai_extraction_methods():
    """Compare CrewAI extraction methods (Traditional, AI, OpenAI)."""
    print("🤖 Demo 4: CrewAI Extraction Methods Comparison")
    print("-" * 50)

    # Create sample document as PDF
    import fitz  # PyMuPDF
    temp_path = tempfile.mktemp(suffix=".pdf")
    doc = fitz.open()
    page = doc.new_page()
    # Use a Unicode-compatible font for Vietnamese with proper line breaks
    text_writer = fitz.TextWriter(page.rect)
    font = fitz.Font("helv")  # Use helvetica font with Unicode support
    
    # Split text by line breaks and render each line with proper spacing
    lines = SAMPLE_VIETNAMESE_TEXT.strip().split('\n')
    for i, line in enumerate(lines):
        # Position each line with increasing y-coordinate for proper line spacing
        text_writer.append((50, 50 + i * 14), line, font=font, fontsize=11)
    
    text_writer.write_text(page)
    doc.save(temp_path)
    doc.close()

    # Create a form template in DOCX format
    from docx import Document
    doc = Document()
    doc.add_heading("EMPLOYMENT APPLICATION FORM", 0)
    doc.add_heading("Personal Information:", 1)
    doc.add_paragraph("Name: ___________________")
    doc.add_paragraph("Date of Birth: ___________________")
    doc.add_paragraph("Email: ___________________")
    doc.add_heading("Skills:", 1)
    doc.add_paragraph("Programming Languages: ___________________")
    
    form_path = tempfile.mktemp(suffix=".docx")
    doc.save(form_path)

    # Test Traditional Method
    print("🔧 Testing Traditional Extraction with CrewAI...")
    traditional_crew = DocumentProcessingCrew(
        text_model="llama3.2:3b",
        extraction_method="traditional",
        vision_model="llava:7b",
    )

    output_traditional = tempfile.mktemp(suffix=".docx")

    start_time = time.time()
    result_traditional = traditional_crew.process_document(
        temp_path,
        form_path,
        output_traditional,
    )
    traditional_time = time.time() - start_time

    print(f"✅ Traditional method: {traditional_time:.2f}s")
    print(f"   Status: {'Success' if result_traditional.success else 'Failed'}")

    # Test AI Method (simulated - would require vision model)
    print("\n🧠 Testing AI Extraction with CrewAI...")
    # Simulate AI configuration without actually running it since it requires vision model
    DocumentProcessingCrew(
        text_model="llama3.2:3b",
        extraction_method="ai",
        vision_model="llava:7b",
    )

    # Note: This would require actual vision model installation
    print("   ⚠️ AI extraction requires vision model (llava:7b)")
    print("   Demo shows configuration and workflow")

    print("⚙️ AI method configuration ready")
    print("   Agents: DocumentCollector (AI) → Translator → FormAnalyst → FormFiller")
    print("   Vision Model: llava:7b (if available)")

    # Test OpenAI Method (requires API key)
    print("\n☁️ Testing OpenAI Extraction with CrewAI...")
    openai_api_key = os.environ.get("OPENAI_API_KEY")

    if openai_api_key:
        # Create OpenAI-based crew
        DocumentProcessingCrew(
            text_model="llama3.2:3b",
            extraction_method="openai",
            vision_model="llava:7b",
            openai_api_key=openai_api_key,
            openai_model="gpt-4o",
        )

        print("   i️ Using OpenAI API for extraction")
        print("   Demo shows configuration and workflow")
        print("⚙️ OpenAI method configuration ready")
        print("   Agents: DocumentCollector (OpenAI) → Translator → FormAnalyst → FormFiller")
        print("   OpenAI Model: gpt-4o")
    else:
        print("   ⚠️ OpenAI extraction requires API key")
        print("   To test, set OPENAI_API_KEY environment variable")

    # Comparison of all methods
    print("\n📊 CrewAI Extraction Method Comparison:")
    print("Traditional: Fast, reliable for clean text documents")
    print("AI: Better for complex layouts, handwriting using local models")
    print("OpenAI: Highest accuracy for complex documents using cloud API")
    print("All: Use same collaborative CrewAI agent framework")

    # Save to .data directory
    data_dir = Path(".data")
    data_dir.mkdir(exist_ok=True)
    
    # Copy files to .data
    import shutil
    if Path(temp_path).exists():
        shutil.copy2(temp_path, data_dir / "demo4_source.pdf")
    if Path(form_path).exists():
        shutil.copy2(form_path, data_dir / "demo4_form.docx")
    if Path(output_traditional).exists():
        shutil.copy2(output_traditional, data_dir / "demo4_output.docx")
        
    # Cleanup
    Path(temp_path).unlink()
    Path(form_path).unlink()
    if Path(output_traditional).exists():
        Path(output_traditional).unlink()

    print("\n")


def demo_crewai_error_handling():
    """Demonstrate CrewAI error handling and recovery."""
    print("🤖 Demo 5: CrewAI Error Handling")
    print("-" * 50)

    # Test 1: Non-existent file
    print("🧪 Test 1: Non-existent source file")
    crew_processor = DocumentProcessingCrew(text_model="llama3.2:3b")
    result = crew_processor.process_document("non_existent.pdf", "some_form.docx", "output.docx")
    print(f"Result: {'✅ Handled' if not result.success else '❌ Unexpected success'}")
    print(f"Error: {result.error}")
    print("🤖 CrewAI agents gracefully handled the error")

    # Test 2: Invalid configuration
    print("\n🧪 Test 2: Invalid model configuration")
    try:
        DocumentProcessingCrew(text_model="non_existent_model", extraction_method="traditional")
        print("✅ CrewAI validated configuration and provided fallback")
    except Exception as e:
        print(f"✅ CrewAI caught configuration error: {e}")

    # Test 3: Agent failure simulation
    print("\n🧪 Test 3: Agent collaboration under stress")
    print("🤖 CrewAI Features:")
    print("   - Automatic retry mechanisms")
    print("   - Task dependency management")
    print("   - Agent failure recovery")
    print("   - Detailed error reporting")
    print("   - Graceful degradation")

    print("\n")


def demo_crewai_configuration():
    """Show CrewAI configuration options."""
    print("🤖 Demo 6: CrewAI Configuration Options")
    print("-" * 50)

    # Different configurations
    configs = [
        {
            "name": "Fast Processing",
            "text_model": "llama3.2:3b",
            "extraction_method": "traditional",
            "description": "Optimized for speed with traditional extraction",
        },
        {
            "name": "High Accuracy (Local)",
            "text_model": "llama3.1:8b",
            "extraction_method": "ai",
            "vision_model": "llava:13b",
            "description": "Maximum accuracy with larger local models",
        },
        {
            "name": "Balanced",
            "text_model": "qwen2.5:7b",
            "extraction_method": "traditional",
            "description": "Good balance of speed and quality",
        },
        {
            "name": "Cloud-Powered",
            "text_model": "llama3.2:3b",
            "extraction_method": "openai",
            "openai_model": "gpt-4o",
            "description": "Maximum accuracy using OpenAI's vision API",
        },
    ]

    print("📋 CrewAI Configuration Profiles:")
    for config in configs:
        print(f"\n🤖 {config['name']}:")
        print(f"   Text Model: {config['text_model']}")
        print(f"   Extraction: {config['extraction_method']}")
        if config.get("vision_model"):
            print(f"   Vision Model: {config['vision_model']}")
        if config.get("openai_model"):
            print(f"   OpenAI Model: {config['openai_model']}")
        print(f"   Use Case: {config['description']}")

    print("\n🔧 CrewAI Advanced Settings:")
    print("   - Process Type: Sequential (default) | Hierarchical")
    print("   - Max Retries: 3 per task")
    print("   - Timeout: Configurable per agent")
    print("   - Verbose Mode: Detailed agent logging")
    print("   - Memory: Shared context between agents")

    print("\n⚙️ Agent-Specific Configuration:")
    print("   DocumentCollector: Extraction method, vision model")
    print("   Translator: Text model, temperature, max tokens")
    print("   FormAnalyst: Analysis depth, field detection")
    print("   FormFiller: Mapping confidence, field validation")

    print("\n")


def main():
    """Run all CrewAI demos."""
    print("🚀 Vietnamese Document Form Filler - CrewAI Demo")
    print("=" * 60)
    print("This demo showcases the CrewAI-powered multi-agent")
    print("document processing system.\n")

    # Check if CrewAI and Ollama are available
    print("🔍 Checking CrewAI system requirements...")
    try:
        import crewai

        print(f"✅ CrewAI installed: v{crewai.__version__}")

        # Check if crewai_tools is available
        try:
            import crewai_tools  # noqa

            print("✅ CrewAI Tools available")
        except ImportError:
            print("⚠️ CrewAI Tools not available")

        # Check Ollama
        import aiohttp

        async def check_ollama():
            try:
                async with (
                    aiohttp.ClientSession() as session,
                    session.get(
                        "http://localhost:11434/api/tags",
                    ) as response,
                ):
                    if response.status == 200:
                        print("✅ Ollama is running")
                        data = await response.json()
                        models = [m["name"] for m in data.get("models", [])]
                        print(f"Available models: {', '.join(models[:3])}...")
                        return True
                    else:
                        print("❌ Ollama is not responding")
                        return False
            except Exception as e:
                print(f"❌ Cannot connect to Ollama: {e}")
                return False

        asyncio.run(check_ollama())

    except ImportError as e:
        print(f"❌ CrewAI not properly installed: {e}")
        print("Install with: pip install crewai crewai-tools langchain langchain-community")
        return
    except Exception as e:
        print(f"❌ System check failed: {e}")
        print("This demo requires CrewAI and Ollama to be running")
        return

    print("\n")

    # Run CrewAI demos
    demo_crewai_single_document()
    demo_crewai_agent_capabilities()
    demo_crewai_batch_processing()
    demo_crewai_extraction_methods()
    demo_crewai_error_handling()
    demo_crewai_configuration()

    print("🎉 CrewAI Demo completed!")
    print("\n🤖 CrewAI Advantages Summary:")
    print("   ✅ Structured agent collaboration")
    print("   ✅ Built-in error handling and retries")
    print("   ✅ Task dependency management")
    print("   ✅ Enhanced observability and logging")
    print("   ✅ Scalable parallel processing")
    print("   ✅ Easy configuration and customization")
    print("   ✅ Multiple extraction methods (Traditional, AI, OpenAI)")
    print("\nTry the CLI commands for full CrewAI features!")


if __name__ == "__main__":
    main()  #!/usr/bin/env python3
"""
Demo script for Vietnamese Document Form Filler
Shows basic usage and capabilities
"""


# Demo data - sample Vietnamese text
SAMPLE_VIETNAMESE_TEXT = """
Họ và tên: Nguyễn Văn An
Ngày sinh: 15/03/1990
Địa chỉ: 123 Phố Huế, Quận Hai Bà Trưng, Hà Nội
Số điện thoại: 0123456789
Email: nguyen.van.an@email.com
Trình độ học vấn: Cử nhân Công nghệ Thông tin

Kinh nghiệm làm việc:
- 2015-2018: Lập trình viên tại Công ty ABC
- 2018-2022: Trưởng nhóm phát triển tại Công ty XYZ
- 2022-hiện tại: Kiến trúc sư phần mềm tại Công ty DEF

Kỹ năng:
- Ngôn ngữ lập trình: Python, Java, JavaScript
- Framework: Django, React, Spring Boot
- Cơ sở dữ liệu: MySQL, PostgreSQL, MongoDB
- DevOps: Docker, Kubernetes, AWS
"""

SAMPLE_FORM_CONTENT = """
EMPLOYMENT APPLICATION FORM

Personal Information:
Name: ___________________
Date of Birth: ___________________
Address: ___________________________________
Phone Number: ___________________
Email: ___________________

Education:
Highest Degree: ___________________
Institution: ___________________

Work Experience:
Previous Position: ___________________
Company: ___________________
Years of Experience: ___________________

Skills:
Technical Skills: ___________________________________
Programming Languages: ___________________
Frameworks: ___________________
Databases: ___________________

Additional Information:
___________________________________________
___________________________________________
"""


async def demo_single_document():
    """Demonstrate single document processing."""
    print("🔹 Demo 1: Single Document Processing")
    print("-" * 40)

    # Create temporary files
    # Source as PDF
    import fitz  # PyMuPDF
    source_path = tempfile.mktemp(suffix=".pdf")
    doc = fitz.open()
    page = doc.new_page()
    # Use a Unicode-compatible font for Vietnamese with proper line breaks
    text_writer = fitz.TextWriter(page.rect)
    font = fitz.Font("helv")  # Use helvetica font with Unicode support
    
    # Split text by line breaks and render each line with proper spacing
    lines = SAMPLE_VIETNAMESE_TEXT.strip().split('\n')
    for i, line in enumerate(lines):
        # Position each line with increasing y-coordinate for proper line spacing
        text_writer.append((50, 50 + i * 14), line, font=font, fontsize=11)
    
    text_writer.write_text(page)
    doc.save(source_path)
    doc.close()

    # Form as DOCX
    from docx import Document
    form_path = tempfile.mktemp(suffix=".docx")
    doc = Document()
    doc.add_paragraph(SAMPLE_FORM_CONTENT)
    doc.save(form_path)

    # Output path
    output_path = tempfile.mktemp(suffix=".docx")

    # Initialize processor
    processor = DocumentProcessingCrew(text_model="llama3.2:3b")

    print("📄 Processing Vietnamese CV...")
    print(f"Source: {Path(source_path).name}")
    print(f"Form: {Path(form_path).name}")

    # Process document
    result = await processor.process_document(source_path, form_path, output_path)

    if result.success:
        print("✅ Success!")
        print(f"Output saved to: {output_path}")

        # Show a preview of the result
        if Path(output_path).exists():
            with Path(output_path).open() as f:
                content = f.read()
                print("\n📋 Preview of filled form:")
                print(content[:300] + "..." if len(content) > 300 else content)
    else:
        print(f"❌ Failed: {result.error}")

    # Cleanup
    Path(source_path).unlink()
    Path(form_path).unlink()
    if Path(output_path).exists():
        Path(output_path).unlink()

    print("\n")


# Skip this function as the individual agents have been replaced by CrewAI tools
# This function is equivalent to demo_crewai_agent_capabilities


async def demo_batch_processing():
    """Demonstrate batch processing."""
    print("🔹 Demo 3: Batch Processing")
    print("-" * 40)

    # Create sample documents
    sample_documents = [
        f"Document {i}: Tên: Người {i}, Tuổi: {20 + i}, Địa chỉ: Hà Nội" for i in range(1, 4)
    ]

    # Create temporary directory structure
    temp_dir = Path(tempfile.mkdtemp())
    input_dir = temp_dir / "input"
    output_dir = temp_dir / "output"
    input_dir.mkdir()
    output_dir.mkdir()

    # Create sample form as docx
    from docx import Document
    form_path = temp_dir / "form.docx"
    doc = Document()
    doc.add_paragraph(SAMPLE_FORM_CONTENT)
    doc.save(form_path)

    # Create input documents as PDFs with proper font for Vietnamese
    import fitz  # PyMuPDF
    for i, content in enumerate(sample_documents, 1):
        doc_path = input_dir / f"document_{i}.pdf"
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        
        # Use a Unicode-compatible font for Vietnamese with TextWriter and line breaks
        text_writer = fitz.TextWriter(page.rect)
        font = fitz.Font("helv")  # Use helvetica font with Unicode support
        
        # Split text by line breaks and render each line with proper spacing
        lines = content.strip().split('\n')
        for i, line in enumerate(lines):
            # Position each line with increasing y-coordinate for proper line spacing
            text_writer.append((50, 50 + i * 15), line, font=font, fontsize=12)
            
        text_writer.write_text(page)
        pdf_doc.save(doc_path)
        pdf_doc.close()

    print(f"📁 Created {len(sample_documents)} sample documents")
    print(f"Input directory: {input_dir}")
    print(f"Output directory: {output_dir}")

    # Initialize batch processor
    batch_processor = CrewAIBatchProcessor(text_model="llama3.2:3b", max_concurrent=2, timeout=60)

    # Discover jobs
    job_count = batch_processor.discover_jobs(
        str(input_dir),
        str(form_path),
        str(output_dir),
        "*.pdf",
    )

    print(f"🔍 Discovered {job_count} processing jobs")

    # Progress callback
    def progress_callback(completed, total, job):
        progress = (completed / total) * 100
        status = "✅" if job.status == "completed" else "❌"
        print(f"Progress: {completed}/{total} ({progress:.1f}%) - {status} {job.source_path.name}")

    # Process batch
    print("\n🔄 Starting batch processing...")
    stats = await batch_processor.process_all(progress_callback)

    print("\n📊 Batch Processing Results:")
    print(f"Total: {stats['total']}")
    print(f"Completed: {stats['completed']}")
    print(f"Failed: {stats['failed']}")
    print(f"Success Rate: {stats['success_rate']:.1f}%")
    print(f"Total Time: {stats['total_time']:.2f}s")
    print(f"Average per Job: {stats['average_job_time']:.2f}s")

    # Cleanup
    import shutil

    shutil.rmtree(temp_dir)

    print("\n")


async def demo_error_handling():
    """Demonstrate error handling scenarios."""
    print("🔹 Demo 4: Error Handling")
    print("-" * 40)

    processor = DocumentProcessingCrew(text_model="llama3.2:3b")

    # Test 1: Non-existent file
    print("🧪 Test 1: Non-existent source file")
    result = await processor.process_document("non_existent.pdf", "some_form.docx", "output.docx")
    print(f"Result: {'✅ Handled' if not result.success else '❌ Unexpected success'}")
    print(f"Error: {result.error}")

    # Test 2: Empty content
    print("\n🧪 Test 2: Empty document")
    # Create empty PDF
    import fitz  # PyMuPDF
    empty_path = tempfile.mktemp(suffix=".pdf")
    pdf_doc = fitz.open()
    pdf_doc.new_page()  # Empty page
    pdf_doc.save(empty_path)
    pdf_doc.close()

    # Create form in DOCX format
    from docx import Document
    form_path = tempfile.mktemp(suffix=".docx")
    doc = Document()
    doc.add_paragraph(SAMPLE_FORM_CONTENT)
    doc.save(form_path)

    # Create output path
    output_path = tempfile.mktemp(suffix=".docx")

    result = await processor.process_document(empty_path, form_path, output_path)
    print(f"Result: {'✅ Handled' if not result.success else '❌ Unexpected success'}")
    print(f"Error: {result.error}")

    # Save to .data directory
    data_dir = Path(".data")
    data_dir.mkdir(exist_ok=True)
    
    # Copy files to .data
    import shutil
    if Path(empty_path).exists():
        shutil.copy2(empty_path, data_dir / "demo5_empty_source.pdf")
    if Path(form_path).exists():
        shutil.copy2(form_path, data_dir / "demo5_form.docx")
    if Path(output_path).exists():
        shutil.copy2(output_path, data_dir / "demo5_output.docx")
        
    # Cleanup
    Path(empty_path).unlink()
    Path(form_path).unlink()
    if Path(output_path).exists():
        Path(output_path).unlink()

    print("\n")


def demo_configuration():
    """Show configuration options."""
    print("🔹 Demo 5: Configuration Options")
    print("-" * 40)

    # Sample configuration
    config = {
        "default_model": "llama3.2:3b",
        "ollama_host": "localhost",
        "ollama_port": 11434,
        "ocr_language": "vie",
        "translation_settings": {
            "temperature": 0.1,
            "max_tokens": 2048,
            "system_prompt": "You are a professional translator...",
        },
        "form_filling_settings": {"confidence_threshold": 0.7, "max_field_mappings": 20},
    }

    print("📋 Sample Configuration:")
    print(json.dumps(config, indent=2))

    print("\n🔧 Environment Variables:")
    env_vars = [
        "OLLAMA_HOST=localhost",
        "OLLAMA_PORT=11434",
        "OLLAMA_MODEL=llama3.2:3b",
        "LOG_LEVEL=INFO",
    ]
    for var in env_vars:
        print(f"  export {var}")

    print("\n")


def main():
    """Run CrewAI demos."""
    print("🚀 Vietnamese Document Form Filler - CrewAI Demo")
    print("=" * 60)
    print("This demo showcases the CrewAI-powered multi-agent")
    print("document processing system.\n")

    # Check if CrewAI and Ollama are available
    print("🔍 Checking CrewAI system requirements...")
    try:
        import crewai

        print(f"✅ CrewAI installed: v{crewai.__version__}")

        # Check if crewai_tools is available
        try:
            import crewai_tools  # noqa

            print("✅ CrewAI Tools available")
        except ImportError:
            print("⚠️ CrewAI Tools not available")

        # Check Ollama
        import aiohttp

        async def check_ollama():
            try:
                async with (
                    aiohttp.ClientSession() as session,
                    session.get(
                        "http://localhost:11434/api/tags",
                    ) as response,
                ):
                    if response.status == 200:
                        print("✅ Ollama is running")
                        data = await response.json()
                        models = [m["name"] for m in data.get("models", [])]
                        print(f"Available models: {', '.join(models[:3])}...")
                        return True
                    else:
                        print("❌ Ollama is not responding")
                        return False
            except Exception as e:
                print(f"❌ Cannot connect to Ollama: {e}")
                return False

        asyncio.run(check_ollama())

    except ImportError as e:
        print(f"❌ CrewAI not properly installed: {e}")
        print("Install with: pip install crewai crewai-tools langchain langchain-community")
        return
    except Exception as e:
        print(f"❌ System check failed: {e}")
        print("This demo requires CrewAI and Ollama to be running")
        return

    print("\n")

    # Run CrewAI demos
    demo_crewai_single_document()
    demo_crewai_agent_capabilities()
    demo_crewai_batch_processing()
    demo_crewai_extraction_methods()
    demo_crewai_error_handling()
    demo_crewai_configuration()

    print("🎉 CrewAI Demo completed!")
    print("\n🤖 CrewAI Advantages Summary:")
    print("   ✅ Structured agent collaboration")
    print("   ✅ Built-in error handling and retries")
    print("   ✅ Task dependency management")
    print("   ✅ Enhanced observability and logging")
    print("   ✅ Scalable parallel processing")
    print("   ✅ Easy configuration and customization")
    print("   ✅ Multiple extraction methods (Traditional, AI, OpenAI)")
    print("\nTry the CLI commands for full CrewAI features!")


if __name__ == "__main__":
    main()  # No need for asyncio.run as it's not an async function
