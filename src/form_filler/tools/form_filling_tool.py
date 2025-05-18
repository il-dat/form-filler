#!/usr/bin/env python3
"""
Form Filling Tool for Vietnamese Document Form Filler.

Handles filling DOCX forms with translated content.
"""

import json
import logging
from typing import Any

from crewai.tools import BaseTool
from docx import Document
from pydantic import Field, PrivateAttr

from form_filler.ai_providers import AIProviderFactory
from form_filler.tools.form_analysis_tool import FormAnalysisTool

# Setup logging
logger = logging.getLogger(__name__)


class FormFillingTool(BaseTool):
    """Tool for filling DOCX forms with translated content."""

    name: str = "form_filler"
    description: str = "Fill DOCX form fields with provided content"

    # Provider configuration
    provider_name: str = Field(default="ollama")
    model_name: str = Field(default="llama3.2:3b")
    api_key: str | None = Field(default=None)
    api_base: str | None = Field(default=None)

    # Private attribute for AI provider
    _ai_provider = PrivateAttr(default=None)  # Will hold AIProvider instance

    def __init__(
        self,
        provider_name: str = "ollama",
        model_name: str = "llama3.2:3b",
        api_key: str | None = None,
        api_base: str | None = None,
        *args: Any,
        **kwargs: Any,
    ) -> None:
        """Initialize the form filling tool.

        Args:
            provider_name: The name of the AI provider (ollama, openai, anthropic, etc.)
            model_name: The name of the model to use
            api_key: API key for the provider (if needed)
            api_base: Base URL for the API (if needed)
            *args: Additional positional arguments
            **kwargs: Additional keyword arguments
        """
        # Support for legacy parameters
        if "model" in kwargs and model_name == "llama3.2:3b":
            model_name = kwargs.pop("model")

        # Set parameters
        kwargs["provider_name"] = provider_name
        kwargs["model_name"] = model_name
        kwargs["api_key"] = api_key
        kwargs["api_base"] = api_base

        super().__init__(*args, **kwargs)

        # Initialize AI provider
        try:
            self._ai_provider = AIProviderFactory.create_provider(
                provider_name=self.provider_name,
                model_name=self.model_name,
                api_key=self.api_key,
                api_base=self.api_base,
            )
        except Exception as e:
            logger.error(f"Failed to initialize AI provider: {e}")
            self._ai_provider = None

    def _run(
        self,
        form_path: str,
        translated_text: str,
        output_path: str,
        field_mappings: str | None = None,
    ) -> str:
        """Fill the form with mapped content.

        Args:
            form_path: Path to the form template
            translated_text: Translated content to fill in the form
            output_path: Path to save the filled form
            field_mappings: Optional JSON string with field mappings

        Returns:
            JSON string with information about the filled form

        Raises:
            Exception: If form filling fails
        """
        try:
            doc = Document(form_path)

            # If no field mappings provided, generate them using AI
            if not field_mappings:
                field_mappings = self._generate_field_mappings(form_path, translated_text)

            # Parse field mappings
            try:
                mappings_data = json.loads(field_mappings)
                field_mappings_list = mappings_data.get("field_mappings", [])
            except json.JSONDecodeError:
                # Fallback: create simple mapping
                field_mappings_list = self._create_fallback_mappings(doc, translated_text)

            filled_count = 0

            # Fill paragraphs
            for paragraph in doc.paragraphs:
                for mapping in field_mappings_list:
                    field_text = mapping.get("field_text", "")
                    fill_content = mapping.get("fill_with", "")

                    if field_text in paragraph.text:
                        # Replace placeholder with content
                        new_text = paragraph.text.replace("____", fill_content)
                        new_text = new_text.replace(field_text, fill_content)
                        paragraph.clear()
                        paragraph.add_run(new_text)
                        filled_count += 1

            # Fill tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for mapping in field_mappings_list:
                            field_text = mapping.get("field_text", "")
                            fill_content = mapping.get("fill_with", "")

                            if field_text in cell.text:
                                cell.text = cell.text.replace("____", fill_content)
                                cell.text = cell.text.replace(field_text, fill_content)
                                filled_count += 1

            # Save filled document
            doc.save(output_path)

            return json.dumps(
                {
                    "output_path": output_path,
                    "fields_filled": filled_count,
                    "total_mappings": len(field_mappings_list),
                },
            )

        except Exception as e:
            logger.error(f"Form filling failed: {e}")
            raise

    def _generate_field_mappings(self, form_path: str, content: str) -> str:
        """Generate field mappings using AI.

        Args:
            form_path: Path to the form template
            content: Translated content to fill in the form

        Returns:
            JSON string with field mappings
        """
        # Check if AI provider is initialized
        if not self._ai_provider:
            logger.warning("AI provider not initialized, using fallback mapping")
            form_analyzer = FormAnalysisTool()
            form_fields = form_analyzer._run(form_path)
            return self._create_fallback_json(form_fields, content)

        # Analyze form structure
        form_analyzer = FormAnalysisTool()
        form_fields = form_analyzer._run(form_path)

        system_prompt = """You are an expert at analyzing documents and mapping content to form fields.
        Given a form structure and translated content, determine how to fill each field appropriately.
        Return ONLY valid JSON in the specified format."""

        prompt = f"""Form fields found:
{form_fields}

Content to fill:
{content}

Please analyze and create a mapping of which content should fill which fields.
Consider the context and purpose of each field. Return only valid JSON in this format:
{{
    "field_mappings": [
        {{
            "field_text": "Original field text",
            "fill_with": "Content to fill this field",
            "confidence": 0.95
        }}
    ]
}}"""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ]

        try:
            # Use the chat completion method from our AI provider
            result = self._ai_provider.chat_completion(
                messages=messages,
                max_tokens=2000,  # Larger token limit for complex JSON responses
                temperature=0.2,  # Lower temperature for more consistent JSON
            )

            # Validate JSON in the response
            try:
                # Try to parse the result as JSON
                json.loads(result)
                return result
            except json.JSONDecodeError:
                # If response isn't valid JSON, try to extract JSON part
                if "{" in result and "}" in result:
                    start = result.find("{")
                    end = result.rfind("}") + 1
                    json_str = result[start:end]
                    try:
                        json.loads(json_str)
                        return json_str
                    except json.JSONDecodeError:
                        # If extraction fails, use fallback
                        logger.warning("Failed to extract valid JSON from AI response")
                        return self._create_fallback_json(form_fields, content)
                else:
                    # If no JSON-like content, use fallback
                    logger.warning("AI response didn't contain valid JSON")
                    return self._create_fallback_json(form_fields, content)

        except Exception as e:
            logger.error(f"AI field mapping failed: {e}")
            return self._create_fallback_json(form_fields, content)

    def _create_fallback_mappings(self, doc: Document, content: str) -> list[dict[str, Any]]:
        """Create simple fallback mappings when AI fails."""
        paragraphs_with_placeholders = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and ("____" in text or "[" in text):
                paragraphs_with_placeholders.append(text)

        # Simple mapping: use first part of content for first field, etc.
        content_parts = content.split("\n")[: len(paragraphs_with_placeholders)]

        mappings = []
        for i, field_text in enumerate(paragraphs_with_placeholders):
            if i < len(content_parts):
                mappings.append(
                    {
                        "field_text": field_text,
                        "fill_with": content_parts[i][:100],  # Limit length
                        "confidence": 0.5,
                    },
                )

        return mappings

    def _create_fallback_json(self, form_fields: str, content: str) -> str:
        """Create fallback JSON when AI mapping fails."""
        try:
            fields = json.loads(form_fields)
            content_parts = content.split("\n")

            mappings = []
            for i, field in enumerate(fields[:3]):  # Limit to first 3 fields
                if i < len(content_parts):
                    mappings.append(
                        {
                            "field_text": field.get("text", ""),
                            "fill_with": content_parts[i][:100],
                            "confidence": 0.5,
                        },
                    )

            return json.dumps({"field_mappings": mappings})
        except Exception:
            return json.dumps({"field_mappings": []})
