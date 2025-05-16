#!/usr/bin/env python3
"""
Form Analysis Tool for Vietnamese Document Form Filler.

Handles analyzing DOCX form structure and identifying fillable fields.
"""

import json
import logging

from crewai.tools import BaseTool
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# Setup logging
logger = logging.getLogger(__name__)


class FormAnalysisTool(BaseTool):
    """Tool for analyzing DOCX form structure."""

    name: str = "form_analyzer"
    description: str = "Analyze a DOCX form to identify fillable fields"

    def _run(self, form_path: str) -> str:
        """Analyze the form structure and return field information."""
        try:
            # Get document from path
            doc = Document(form_path)
            form_fields = []

            # Look for placeholder text, form fields, tables
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and ("____" in text or "[" in text or text.endswith(":")):
                    form_fields.append({"type": "paragraph", "text": text, "placeholder": True})

            # Check for tables (common in forms)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and ("____" in cell_text or "[" in cell_text or cell_text.endswith(":")):
                            form_fields.append(
                                {"type": "table_cell", "text": cell_text, "placeholder": True}
                            )

            return json.dumps(form_fields, indent=2)

        except PackageNotFoundError as e:
            logger.error(f"Form analysis failed: {e}")
            # Re-raise as FileNotFoundError for consistent API
            raise FileNotFoundError(f"Form file not found: {form_path}")
        except Exception as e:
            logger.error(f"Form analysis failed: {e}")
            raise
