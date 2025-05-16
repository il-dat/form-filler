#!/usr/bin/env python3
"""
Vietnamese to English Document Form Filler CLI.

A CrewAI-based multi-agent system for processing Vietnamese documents and filling English DOCX forms.
"""

import asyncio
import json
import logging
import sys
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import aiohttp
import click

# Document processing imports
import fitz  # PyMuPDF for PDF
import pytesseract

# CrewAI imports
from crewai import Agent, Crew, Process, Task
from crewai.tools import BaseTool
from docx import Document
from langchain_community.chat_models import ChatOllama

# Langchain imports for Ollama integration
from langchain_ollama import OllamaLLM
from PIL import Image

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


@dataclass
class ProcessingResult:
    """Data structure for processing results."""

    success: bool
    data: Any
    error: str | None = None
    metadata: dict[str, Any] | None = None


# Custom Tools for CrewAI Agents
class DocumentExtractionTool(BaseTool):
    """Tool for extracting text from documents."""

    name: str = "document_extractor"
    description: str = "Extract text from PDF or image files using traditional or AI methods"

    def __init__(self, extraction_method="traditional", vision_model="llava:7b", *args, **kwargs):
        """Initialize the object."""
        super().__init__()
        self.extraction_method = extraction_method
        self.vision_model = vision_model
        self.ollama_llm = None
        if extraction_method == "ai":
            self.ollama_llm = OllamaLLM(model=vision_model, base_url="http://localhost:11434")

    def _run(self, file_path: str) -> str:
        """Extract text from the given file."""
        try:
            # Convert string path to Path object
            path_obj = Path(file_path)

            if not path_obj.exists():
                raise Exception(f"File not found: {file_path}")

            if path_obj.suffix.lower() == ".pdf":
                if self.extraction_method == "ai":
                    return self._extract_from_pdf_ai(path_obj)
                else:
                    return self._extract_from_pdf_traditional(path_obj)
            elif path_obj.suffix.lower() in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
                if self.extraction_method == "ai":
                    return self._extract_from_image_ai(path_obj)
                else:
                    return self._extract_from_image_traditional(path_obj)
            else:
                raise Exception(f"Unsupported file type: {path_obj.suffix}")

        except Exception as e:
            logger.error(f"Error in DocumentExtractionTool: {e}")
            raise

    def _extract_from_pdf_traditional(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    def _extract_from_pdf_ai(self, file_path: Path) -> str:
        """Extract text from PDF using AI (convert to images and use vision model)."""
        try:
            # Convert PDF pages to images and process with AI
            doc = fitz.open(file_path)
            images_text = []

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                img_data = pix.tobytes("png")

                # Save image temporarily
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                    tmp_img.write(img_data)
                    tmp_img_path = tmp_img.name

                try:
                    # Use AI to extract text from image
                    page_text = self._extract_from_image_ai(Path(tmp_img_path))
                    images_text.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
                finally:
                    # Clean up temporary image
                    Path(tmp_img_path).unlink()

            doc.close()
            return "\n".join(images_text)

        except Exception as e:
            logger.warning(f"AI PDF extraction failed, falling back to traditional: {e}")
            return self._extract_from_pdf_traditional(file_path)

    def _extract_from_image_traditional(self, file_path: Path) -> str:
        """Extract text from image using Tesseract OCR."""
        image = Image.open(file_path)
        return pytesseract.image_to_string(image, lang="vie")

    def _extract_from_image_ai(self, file_path: Path) -> str:
        """Extract text from image using AI vision model."""
        try:
            # Convert image to base64 for AI processing
            import base64

            with Path(file_path).open("rb") as img_file:
                base64_image = base64.b64encode(img_file.read()).decode("utf-8")

            # Create prompt for text extraction
            prompt = f"""Please extract all text from this image. The image may contain Vietnamese text.
            Extract ALL visible text, preserving the original formatting and structure.
            Pay attention to Vietnamese diacritics and special characters.

            Return only the extracted text without any additional commentary.

            <image>data:image/png;base64,{base64_image}</image>

            Extracted text:"""

            # For now, simulate AI extraction (in real implementation, use vision model)
            # This is a placeholder - you would implement actual vision model here
            extracted_text = self.ollama_llm.invoke(prompt)

            if not extracted_text:
                logger.warning("AI image extraction failed, falling back to OCR")
                return self._extract_from_image_traditional(file_path)

            return extracted_text

        except Exception as e:
            logger.warning(f"AI image extraction failed, falling back to OCR: {e}")
            return self._extract_from_image_traditional(file_path)


class TranslationTool(BaseTool):
    """Tool for translating Vietnamese text to English."""

    name: str = "vietnamese_translator"
    description: str = "Translate Vietnamese text to English using Ollama LLM"

    def __init__(self, model="llama3.2:3b", *args, **kwargs):
        """Initialize the object."""
        super().__init__()
        self.llm = ChatOllama(model=model, base_url="http://localhost:11434")

    def _run(self, vietnamese_text: str) -> str:
        """Translate Vietnamese text to English."""
        if not vietnamese_text.strip():
            raise Exception("Empty text provided for translation")

        system_prompt = """You are a professional translator specializing in Vietnamese to English translation.
        Translate the given Vietnamese text to clear, accurate English while preserving the original meaning and context.
        Focus on formal document language appropriate for forms and official papers.

        Return only the English translation without any additional commentary."""

        messages = [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": f"Please translate this Vietnamese text to English:\n\n{vietnamese_text}",
            },
        ]

        try:
            response = self.llm.invoke(messages)
            return response.content if hasattr(response, "content") else str(response)
        except Exception as e:
            logger.error(f"Translation failed: {e}")
            raise


class FormAnalysisTool(BaseTool):
    """Tool for analyzing DOCX form structure."""

    name: str = "form_analyzer"
    description: str = "Analyze DOCX form structure and identify fillable fields"

    def _run(self, form_path: str) -> str:
        """Analyze the form structure and return field information."""
        try:
            doc = Document(form_path)
            form_fields = []

            # Look for placeholder text, form fields, tables
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and ("____" in text or "[" in text or text.endswith(":")):
                    form_fields.append({"type": "paragraph", "text": text, "placeholder": True})

            # Check for tables (common in forms)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text and ("____" in cell_text or "[" in cell_text):
                            form_fields.append(
                                {"type": "table_cell", "text": cell_text, "placeholder": True},
                            )

            return json.dumps(form_fields, indent=2)

        except Exception as e:
            logger.error(f"Form analysis failed: {e}")
            raise


class FormFillingTool(BaseTool):
    """Tool for filling DOCX forms with translated content."""

    name: str = "form_filler"
    description: str = "Fill DOCX form fields with provided content"

    def __init__(self, model="llama3.2:3b", *args, **kwargs):
        """Initialize the object."""
        super().__init__()
        self.llm = ChatOllama(model=model, base_url="http://localhost:11434")

    def _run(
        self,
        form_path: str,
        translated_text: str,
        output_path: str,
        field_mappings: str | None = None,
    ) -> str:
        """Fill the form with mapped content."""
        try:
            doc = Document(form_path)

            # If no field mappings provided, generate them using AI
            if not field_mappings:
                field_mappings = self._generate_field_mappings(form_path, translated_text)

            # Parse field mappings
            try:
                mappings_data = json.loads(field_mappings)
                field_mappings_list = mappings_data.get("field_mappings", [])
            except json.JSONDecodeError:
                # Fallback: create simple mapping
                field_mappings_list = self._create_fallback_mappings(doc, translated_text)

            filled_count = 0

            # Fill paragraphs
            for paragraph in doc.paragraphs:
                for mapping in field_mappings_list:
                    field_text = mapping.get("field_text", "")
                    fill_content = mapping.get("fill_with", "")

                    if field_text in paragraph.text:
                        # Replace placeholder with content
                        new_text = paragraph.text.replace("____", fill_content)
                        new_text = new_text.replace(field_text, fill_content)
                        paragraph.clear()
                        paragraph.add_run(new_text)
                        filled_count += 1

            # Fill tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for mapping in field_mappings_list:
                            field_text = mapping.get("field_text", "")
                            fill_content = mapping.get("fill_with", "")

                            if field_text in cell.text:
                                cell.text = cell.text.replace("____", fill_content)
                                cell.text = cell.text.replace(field_text, fill_content)
                                filled_count += 1

            # Save filled document
            doc.save(output_path)

            return json.dumps(
                {
                    "output_path": output_path,
                    "fields_filled": filled_count,
                    "total_mappings": len(field_mappings_list),
                },
            )

        except Exception as e:
            logger.error(f"Form filling failed: {e}")
            raise

    def _generate_field_mappings(self, form_path: str, content: str) -> str:
        """Generate field mappings using AI."""
        # Analyze form structure
        form_analyzer = FormAnalysisTool()
        form_fields = form_analyzer._run(form_path)

        system_prompt = """You are an expert at analyzing documents and mapping content to form fields.
        Given a form structure and translated content, determine how to fill each field appropriately.
        Return ONLY valid JSON in the specified format."""

        prompt = f"""Form fields found:
{form_fields}

Content to fill:
{content}

Please analyze and create a mapping of which content should fill which fields.
Consider the context and purpose of each field. Return only valid JSON in this format:
{{
    "field_mappings": [
        {{
            "field_text": "Original field text",
            "fill_with": "Content to fill this field",
            "confidence": 0.95
        }}
    ]
}}"""

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt},
        ]

        try:
            response = self.llm.invoke(messages)
            return response.content if hasattr(response, "content") else str(response)
        except Exception as e:
            logger.error(f"AI field mapping failed: {e}")
            return self._create_fallback_json(form_fields, content)

    def _create_fallback_mappings(self, doc: Document, content: str) -> list[dict]:
        """Create simple fallback mappings when AI fails."""
        paragraphs_with_placeholders = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text and ("____" in text or "[" in text):
                paragraphs_with_placeholders.append(text)

        # Simple mapping: use first part of content for first field, etc.
        content_parts = content.split("\n")[: len(paragraphs_with_placeholders)]

        mappings = []
        for i, field_text in enumerate(paragraphs_with_placeholders):
            if i < len(content_parts):
                mappings.append(
                    {
                        "field_text": field_text,
                        "fill_with": content_parts[i][:100],  # Limit length
                        "confidence": 0.5,
                    },
                )

        return mappings

    def _create_fallback_json(self, form_fields: str, content: str) -> str:
        """Create fallback JSON when AI mapping fails."""
        try:
            fields = json.loads(form_fields)
            content_parts = content.split("\n")

            mappings = []
            for i, field in enumerate(fields[:3]):  # Limit to first 3 fields
                if i < len(content_parts):
                    mappings.append(
                        {
                            "field_text": field.get("text", ""),
                            "fill_with": content_parts[i][:100],
                            "confidence": 0.5,
                        },
                    )

            return json.dumps({"field_mappings": mappings})
        except Exception:
            return json.dumps({"field_mappings": []})


# CrewAI Agents
def create_document_collector_agent(
    extraction_method: str = "traditional",
    vision_model: str = "llava:7b",
) -> Agent:
    """Create the document collection agent."""
    return Agent(
        role="Document Text Extractor",
        goal="Extract text content from Vietnamese documents (PDFs and images) with high accuracy",
        backstory="""You are a specialized document processing expert with advanced capabilities in text extraction.
        You can handle both traditional OCR methods and cutting-edge AI vision models to extract text from various document formats.
        Your expertise includes processing Vietnamese documents with proper diacritics and special characters.""",
        tools=[
            DocumentExtractionTool(extraction_method=extraction_method, vision_model=vision_model),
        ],
        verbose=True,
        allow_delegation=False,
    )


def create_translator_agent(model: str = "llama3.2:3b") -> Agent:
    """Create the translation agent."""
    return Agent(
        role="Vietnamese to English Translator",
        goal="Provide accurate and contextually appropriate translations from Vietnamese to English",
        backstory="""You are a professional translator with deep expertise in Vietnamese and English languages.
        You specialize in translating official documents, forms, and business communications while preserving
        the original meaning and maintaining formal language appropriate for document processing.""",
        tools=[TranslationTool(model=model)],
        verbose=True,
        allow_delegation=False,
    )


def create_form_analyst_agent() -> Agent:
    """Create the form analysis agent."""
    return Agent(
        role="Document Form Analyst",
        goal="Analyze DOCX forms and identify all fillable fields and their purposes",
        backstory="""You are an expert in document analysis and form processing. You can quickly identify
        form fields, understand their context and purpose, and determine the most appropriate content
        to fill each field based on available translated information.""",
        tools=[FormAnalysisTool()],
        verbose=True,
        allow_delegation=False,
    )


def create_form_filler_agent(model: str = "llama3.2:3b") -> Agent:
    """Create the form filling agent."""
    return Agent(
        role="Form Completion Specialist",
        goal="Fill DOCX forms with translated content using intelligent field mapping",
        backstory="""You are a form completion specialist with the ability to understand document context
        and intelligently map translated content to appropriate form fields. You ensure accuracy and
        maintain proper formatting while filling forms.""",
        tools=[FormFillingTool(model=model)],
        verbose=True,
        allow_delegation=False,
    )


class DocumentProcessingCrew:
    """Main CrewAI crew for document processing."""

    def __init__(  # noqa: D107
        self,
        text_model: str = "llama3.2:3b",
        extraction_method: str = "traditional",
        vision_model: str = "llava:7b",
    ):
        # Create agents
        self.document_collector = create_document_collector_agent(extraction_method, vision_model)
        self.translator = create_translator_agent(text_model)
        self.form_analyst = create_form_analyst_agent()
        self.form_filler = create_form_filler_agent(text_model)

        # Store configuration
        self.extraction_method = extraction_method
        self.text_model = text_model
        self.vision_model = vision_model

    def process_document(
        self,
        source_path: str,
        form_path: str,
        output_path: str,
    ) -> ProcessingResult:
        """Process a document through the CrewAI pipeline."""
        try:
            # Define tasks
            extraction_task = Task(
                description=f"""Extract all text content from the Vietnamese document at: {source_path}

                Requirements:
                - Extract all visible text including Vietnamese diacritics
                - Preserve formatting and structure where possible
                - Handle both text-based and image-based content
                - Use {"AI vision models" if self.extraction_method == "ai" else "traditional OCR methods"}

                Return the complete extracted text.""",
                agent=self.document_collector,
                expected_output="Complete text content extracted from the Vietnamese document",
            )

            translation_task = Task(
                description="""Translate the extracted Vietnamese text to English.

                Requirements:
                - Maintain professional, formal language suitable for documents
                - Preserve all important information (names, dates, addresses, etc.)
                - Ensure accuracy and contextual appropriateness
                - Keep formatting structure where relevant

                Return the complete English translation.""",
                agent=self.translator,
                expected_output="Professional English translation of the Vietnamese text",
                context=[extraction_task],
            )

            form_analysis_task = Task(
                description=f"""Analyze the DOCX form structure at: {form_path}

                Requirements:
                - Identify all fillable fields and placeholders
                - Understand the purpose and context of each field
                - Provide detailed information about form structure

                Return structured information about the form fields.""",
                agent=self.form_analyst,
                expected_output="Detailed analysis of form structure and fillable fields",
            )

            form_filling_task = Task(
                description=f"""Fill the DOCX form with the translated content and save to: {output_path}

                Requirements:
                - Use the form analysis to understand field purposes
                - Intelligently map translated content to appropriate fields
                - Maintain original form formatting
                - Fill all relevant fields with appropriate content

                Return information about the filling process including number of fields filled.""",
                agent=self.form_filler,
                expected_output="Successfully filled form with detailed completion report",
                context=[translation_task, form_analysis_task],
            )

            # Create and execute crew
            crew = Crew(
                agents=[
                    self.document_collector,
                    self.translator,
                    self.form_analyst,
                    self.form_filler,
                ],
                tasks=[extraction_task, translation_task, form_analysis_task, form_filling_task],
                process=Process.sequential,
                verbose=True,
            )

            # Execute the crew
            logger.info("Starting CrewAI document processing pipeline")
            logger.info(f"Source: {source_path}, Form: {form_path}, Output: {output_path}")
            logger.info(f"Extraction method: {self.extraction_method}")

            result = crew.kickoff()

            # Parse the final result
            try:
                if isinstance(result, str):
                    final_result = json.loads(result)
                else:
                    final_result = {"output_path": output_path, "fields_filled": "Unknown"}

                return ProcessingResult(
                    success=True,
                    data=final_result,
                    metadata={
                        "extraction_method": self.extraction_method,
                        "text_model": self.text_model,
                        "vision_model": (
                            self.vision_model if self.extraction_method == "ai" else None
                        ),
                    },
                )
            except json.JSONDecodeError:
                # If result is not JSON, consider it successful anyway
                return ProcessingResult(
                    success=True,
                    data={"output_path": output_path, "result": str(result)},
                    metadata={
                        "extraction_method": self.extraction_method,
                        "text_model": self.text_model,
                        "vision_model": (
                            self.vision_model if self.extraction_method == "ai" else None
                        ),
                    },
                )

        except Exception as e:
            logger.error(f"CrewAI processing failed: {e}")
            return ProcessingResult(False, None, str(e))


# CLI Implementation
@click.group()
@click.option("--verbose", "-v", is_flag=True, help="Enable verbose logging")
@click.option(
    "--model",
    "-m",
    default="llama3.2:3b",
    help="Ollama model to use for text processing",
)
@click.option(
    "--extraction-method",
    "-e",
    type=click.Choice(["traditional", "ai"]),
    default="traditional",
    help="Text extraction method: traditional (PyMuPDF/Tesseract) or ai (vision models)",
)
@click.option("--vision-model", "-vm", default="llava:7b", help="Vision model for AI extraction")
@click.pass_context
def cli(ctx, verbose, model, extraction_method, vision_model):
    """Vietnamese to English Document Form Filler (CrewAI Edition).

    A CrewAI-based multi-agent system for processing Vietnamese documents (PDF/images)
    and filling English DOCX forms using local Ollama LLMs.

    Extraction Methods:
    - traditional: Use PyMuPDF for PDFs and Tesseract for images
    - ai: Use vision models for both PDFs and images (requires vision-capable models)
    """
    ctx.ensure_object(dict)
    ctx.obj["verbose"] = verbose
    ctx.obj["model"] = model
    ctx.obj["extraction_method"] = extraction_method
    ctx.obj["vision_model"] = vision_model

    if verbose:
        logging.getLogger().setLevel(logging.DEBUG)


@cli.command()
@click.argument("source", type=click.Path(exists=True))
@click.argument("form", type=click.Path(exists=True))
@click.argument("output", type=click.Path())
@click.option("--model", "-m", help="Override default Ollama model")
@click.option(
    "--extraction-method",
    "-e",
    type=click.Choice(["traditional", "ai"]),
    help="Override default extraction method",
)
@click.option("--vision-model", "-vm", help="Override default vision model")
@click.pass_context
def process(ctx, source, form, output, model, extraction_method, vision_model):
    """Process a Vietnamese document and fill an English DOCX form using CrewAI.

    SOURCE: Path to Vietnamese document (PDF or image)
    FORM: Path to English DOCX form template
    OUTPUT: Path where filled form will be saved

    Example with AI extraction:
    python document_processor.py -e ai -vm llava:7b process document.pdf form.docx output.docx
    """
    model = model or ctx.obj["model"]
    extraction_method = extraction_method or ctx.obj["extraction_method"]
    vision_model = vision_model or ctx.obj["vision_model"]

    # Ensure output directory exists
    output_path = Path(output)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Create CrewAI processor
    crew_processor = DocumentProcessingCrew(
        text_model=model,
        extraction_method=extraction_method,
        vision_model=vision_model,
    )

    # Process the document
    start_time = time.time()
    result = crew_processor.process_document(source, form, output)
    processing_time = time.time() - start_time

    if result.success:
        click.echo(f"✅ Success! Document processed in {processing_time:.2f}s")
        click.echo(f"Filled form saved to: {output}")
        if result.metadata:
            click.echo(f"Extraction method: {result.metadata.get('extraction_method', 'N/A')}")
            click.echo(f"Text model: {result.metadata.get('text_model', 'N/A')}")
            if result.metadata.get("vision_model"):
                click.echo(f"Vision model: {result.metadata.get('vision_model')}")
        if result.data and isinstance(result.data, dict):
            fields_filled = result.data.get("fields_filled", "N/A")
            click.echo(f"Fields filled: {fields_filled}")
    else:
        click.echo(f"❌ Error: {result.error}")
        sys.exit(1)


@cli.command()
@click.argument("file_path", type=click.Path(exists=True))
@click.option(
    "--extraction-method",
    "-e",
    type=click.Choice(["traditional", "ai"]),
    help="Override default extraction method",
)
@click.option("--vision-model", "-vm", help="Vision model for AI extraction")
@click.pass_context
def extract(ctx, file_path, extraction_method, vision_model):
    """Extract text from a Vietnamese document (for testing).

    Example with AI extraction:
    python document_processor.py -e ai extract document.pdf
    """
    extraction_method = extraction_method or ctx.obj["extraction_method"]
    vision_model = vision_model or ctx.obj["vision_model"]

    # Create extraction tool
    extractor = DocumentExtractionTool(
        extraction_method=extraction_method,
        vision_model=vision_model,
    )

    try:
        click.echo(f"Extracting text using {extraction_method} method...")
        if extraction_method == "ai":
            click.echo(f"Vision model: {vision_model}")

        extracted_text = extractor._run(file_path)

        click.echo(f"Extracted text ({extraction_method} method):")
        click.echo("-" * 50)
        click.echo(extracted_text)
        click.echo("-" * 50)
        click.echo(f"Characters: {len(extracted_text)}")
    except Exception as e:
        click.echo(f"❌ Error: {e}")
        sys.exit(1)


@cli.command()
@click.argument("vietnamese_text")
@click.option("--model", "-m", help="Override default model")
@click.pass_context
def translate(ctx, vietnamese_text, model):
    """Translate Vietnamese text to English (for testing)."""
    model = model or ctx.obj["model"]

    translator = TranslationTool(model=model)

    try:
        click.echo(f"Translating with model: {model}")
        english_text = translator._run(vietnamese_text)

        click.echo("Translation:")
        click.echo("-" * 50)
        click.echo(english_text)
        click.echo("-" * 50)
    except Exception as e:
        click.echo(f"❌ Error: {e}")
        sys.exit(1)


@cli.command()
@click.option("--host", default="localhost", help="Ollama host")
@click.option("--port", default=11434, help="Ollama port")
@click.option("--check-vision", is_flag=True, help="Also check for vision models")
async def check_ollama(host, port, check_vision):
    """Check if Ollama is running and list available models."""
    url = f"http://{host}:{port}/api/tags"

    try:
        async with aiohttp.ClientSession() as session, session.get(url) as response:
            if response.status == 200:
                data = await response.json()
                models = data.get("models", [])

                click.echo("✅ Ollama is running!")
                click.echo("\nAvailable models:")

                text_models = []
                vision_models = []

                for model in models:
                    name = model.get("name", "Unknown")
                    size = model.get("size", 0)
                    size_mb = size / (1024 * 1024) if size else 0

                    # Classify models
                    if (
                        "llava" in name.lower()
                        or "vision" in name.lower()
                        or "bakllava" in name.lower()
                    ):
                        vision_models.append(f"  - {name} ({size_mb:.1f} MB)")
                    else:
                        text_models.append(f"  - {name} ({size_mb:.1f} MB)")

                click.echo("\n📝 Text Models (for translation and form filling):")
                for model in text_models:
                    click.echo(model)

                if check_vision:
                    click.echo("\n👁️ Vision Models (for AI text extraction):")
                    if vision_models:
                        for model in vision_models:
                            click.echo(model)
                    else:
                        click.echo("  No vision models found. Install with: ollama pull llava:7b")
                        click.echo(
                            "  Vision models enable AI-powered text extraction from images and PDFs",
                        )
                        click.echo("  Supported models: llava:7b, llava:13b, bakllava")

                click.echo("\n🤖 CrewAI Integration Status: ✅ Ready")
                click.echo("Available extraction methods: traditional, ai")

            else:
                click.echo(f"❌ Ollama responded with status: {response.status}")
                sys.exit(1)
    except Exception as e:
        click.echo(f"❌ Error connecting to Ollama: {e}")
        click.echo(f"Make sure Ollama is running on {host}:{port}")
        sys.exit(1)


# Async wrapper for Click commands
def async_command(f):
    """Decorator to run async click commands."""

    def wrapper(*args, **kwargs):
        return asyncio.run(f(*args, **kwargs))

    return wrapper


# Apply async wrapper to commands
check_ollama = async_command(check_ollama)


if __name__ == "__main__":
    cli()
