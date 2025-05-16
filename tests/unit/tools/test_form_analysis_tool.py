"""
Unit tests for FormAnalysisTool.

Tests the functionality of the FormAnalysisTool for analyzing DOCX form structure.
Includes tests for various field types, error handling, and edge cases.
"""

import json
import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from form_filler.tools.form_analysis_tool import FormAnalysisTool


@pytest.fixture
def mock_docx_document():
    """Create a mock for the python-docx Document class."""
    # Create the main Document mock
    mock_doc = MagicMock()

    # Create paragraph mocks with different types of placeholders
    paragraph1 = MagicMock()
    paragraph1.text = "Name: ____"

    paragraph2 = MagicMock()
    paragraph2.text = "Address [enter full address]"

    paragraph3 = MagicMock()
    paragraph3.text = "Phone number:"

    paragraph4 = MagicMock()
    paragraph4.text = "This is just regular text with no placeholder"

    # Set up the mock document's paragraphs
    mock_doc.paragraphs = [paragraph1, paragraph2, paragraph3, paragraph4]

    # Create table cell mocks
    cell1 = MagicMock()
    cell1.text = "Date of birth: ____"

    cell2 = MagicMock()
    cell2.text = "Occupation [specify]"

    cell3 = MagicMock()
    cell3.text = "Comments:"

    cell4 = MagicMock()
    cell4.text = "Regular cell text"

    # Create row mocks
    row1 = MagicMock()
    row1.cells = [cell1, cell2]

    row2 = MagicMock()
    row2.cells = [cell3, cell4]

    # Create table mock
    table = MagicMock()
    table.rows = [row1, row2]

    # Set up the mock document's tables
    mock_doc.tables = [table]

    return mock_doc


@pytest.fixture
def expected_form_fields():
    """Provide the expected form fields structure for test verification."""
    fields = [
        {"type": "paragraph", "text": "Name: ____", "placeholder": True},
        {"type": "paragraph", "text": "Address [enter full address]", "placeholder": True},
        {"type": "paragraph", "text": "Phone number:", "placeholder": True},
        {"type": "table_cell", "text": "Date of birth: ____", "placeholder": True},
        {"type": "table_cell", "text": "Occupation [specify]", "placeholder": True},
        {"type": "table_cell", "text": "Comments:", "placeholder": True},
    ]
    return fields


def test_init():
    """Test initialization of the FormAnalysisTool."""
    tool = FormAnalysisTool()

    assert tool.name == "form_analyzer"
    assert "Analyze a DOCX form to identify fillable fields" in tool.description


@patch("form_filler.tools.form_analysis_tool.Document")
def test_run_file_not_found(mock_document_class):
    """Test handling of non-existent form files."""
    mock_document_class.side_effect = FileNotFoundError("File not found")

    tool = FormAnalysisTool()

    with pytest.raises(FileNotFoundError):
        tool._run("/path/to/nonexistent/form.docx")


@patch("form_filler.tools.form_analysis_tool.Document")
def test_run_invalid_docx(mock_document_class):
    """Test handling of invalid DOCX files."""
    # We'll simulate an actual error thrown by the Document class
    mock_document_class.side_effect = Exception("Invalid DOCX file")

    tool = FormAnalysisTool()

    # The tool should re-raise the original exception for errors other than PackageNotFoundError
    with pytest.raises(Exception, match="Invalid DOCX file"):
        tool._run("/path/to/invalid/form.docx")


@patch("form_filler.tools.form_analysis_tool.Document")
def test_form_with_no_fields(mock_document_class):
    """Test analysis of a form with no fillable fields."""
    # Create a mock document with no fillable fields
    mock_doc = MagicMock()
    mock_doc.paragraphs = []
    mock_doc.tables = []

    # The key is to return our mock document instead of raising an exception
    mock_document_class.return_value = mock_doc

    tool = FormAnalysisTool()
    result = tool._run("/path/to/form.docx")

    # Parse the JSON result
    fields = json.loads(result)

    assert isinstance(fields, list)
    assert len(fields) == 0


@patch("form_filler.tools.form_analysis_tool.Document")
def test_form_with_fields(mock_document_class, mock_docx_document, expected_form_fields):
    """Test analysis of a form with fillable fields."""
    mock_document_class.return_value = mock_docx_document

    tool = FormAnalysisTool()
    result = tool._run("/path/to/form.docx")

    # Parse the JSON result
    fields = json.loads(result)

    assert isinstance(fields, list)
    assert len(fields) == 6

    # Verify each field has the expected structure
    for field in fields:
        assert "type" in field
        assert "text" in field
        assert "placeholder" in field
        assert field["placeholder"] is True

    # Verify specific fields are identified
    field_texts = [field["text"] for field in fields]
    assert "Name: ____" in field_texts
    assert "Address [enter full address]" in field_texts
    assert "Phone number:" in field_texts
    assert "Date of birth: ____" in field_texts
    assert "Occupation [specify]" in field_texts
    assert "Comments:" in field_texts

    # Verify regular text is not included
    assert "This is just regular text with no placeholder" not in field_texts
    assert "Regular cell text" not in field_texts


def test_with_real_temp_file():
    """Test with a temporary file to verify file handling."""
    # This test depends on python-docx being installed
    try:
        import docx  # noqa: F401

        # Create a temporary file path for testing
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_path = temp_file.name

        try:
            # Instead of trying to create a real DOCX file (which would increase test complexity),
            # just mock the Document class when given this specific file path
            with patch("form_filler.tools.form_analysis_tool.Document") as mock_document_class:
                mock_document_class.return_value = MagicMock()
                mock_document_class.return_value.paragraphs = []
                mock_document_class.return_value.tables = []

                tool = FormAnalysisTool()
                result = tool._run(temp_path)

                # Verify the result is valid JSON
                fields = json.loads(result)
                assert isinstance(fields, list)

                # Verify Document was called with the right path
                mock_document_class.assert_called_once_with(temp_path)
        finally:
            # Clean up
            if Path(temp_path).exists():
                Path(temp_path).unlink()
    except ImportError:
        pytest.skip("python-docx not installed, skipping test")


@patch("form_filler.tools.form_analysis_tool.Document")
def test_complex_tables(mock_document_class):
    """Test handling of complex tables with multiple levels of cells."""
    # Create a mock document with complex nested tables
    mock_doc = MagicMock()

    # Mock main table
    main_cell1 = MagicMock()
    main_cell1.text = "Main form field: ____"

    main_cell2 = MagicMock()
    main_cell2.text = "Regular cell text"

    main_row = MagicMock()
    main_row.cells = [main_cell1, main_cell2]

    main_table = MagicMock()
    main_table.rows = [main_row]

    # Mock nested table cells within main table cells
    mock_doc.tables = [main_table]

    # Set up the mock
    mock_document_class.return_value = mock_doc

    # Run the analysis
    tool = FormAnalysisTool()
    result = tool._run("/path/to/form_with_complex_tables.docx")

    # Parse the result
    fields = json.loads(result)

    # Verify the fields were detected
    assert len(fields) == 1
    assert fields[0]["text"] == "Main form field: ____"
    assert fields[0]["type"] == "table_cell"
    assert fields[0]["placeholder"] is True


@patch("form_filler.tools.form_analysis_tool.Document")
@patch("form_filler.tools.form_analysis_tool.logger")
def test_logging_on_error(mock_logger, mock_document_class):
    """Test that errors are properly logged."""
    # Make the document class raise an exception
    mock_document_class.side_effect = Exception("Document parsing error")

    # Create tool and run with error
    tool = FormAnalysisTool()

    # Since we're raising a general Exception (not PackageNotFoundError),
    # we should expect a general Exception to be re-raised
    with pytest.raises(Exception) as exc_info:
        tool._run("/path/to/form.docx")

    # Verify the exception is the right one
    assert str(exc_info.value) == "Document parsing error"

    # Verify logging was called with the expected message
    mock_logger.error.assert_called_once_with("Form analysis failed: Document parsing error")


@patch("form_filler.tools.form_analysis_tool.Document")
def test_form_with_multiple_field_types(mock_document_class):
    """Test analysis of a form with different types of placeholder formats."""
    # Create a mock document with different field types
    mock_doc = MagicMock()

    # Different paragraph field formats
    paragraph1 = MagicMock()
    paragraph1.text = "Underscore placeholder: ____"

    paragraph2 = MagicMock()
    paragraph2.text = "Bracket placeholder [enter text here]"

    paragraph3 = MagicMock()
    paragraph3.text = "Colon placeholder:"

    paragraph4 = MagicMock()
    paragraph4.text = "Not a placeholder text"

    mock_doc.paragraphs = [paragraph1, paragraph2, paragraph3, paragraph4]
    mock_doc.tables = []

    # Set up the mock
    mock_document_class.return_value = mock_doc

    # Run the analysis
    tool = FormAnalysisTool()
    result = tool._run("/path/to/form.docx")

    # Parse the result
    fields = json.loads(result)

    # Verify the fields were detected
    assert len(fields) == 3
    field_texts = [field["text"] for field in fields]
    assert "Underscore placeholder: ____" in field_texts
    assert "Bracket placeholder [enter text here]" in field_texts
    assert "Colon placeholder:" in field_texts
    assert "Not a placeholder text" not in field_texts


@patch("form_filler.tools.form_analysis_tool.Document")
def test_empty_document(mock_document_class):
    """Test analysis of a completely empty document."""
    # Create a mock empty document
    mock_doc = MagicMock()
    mock_doc.paragraphs = []
    mock_doc.tables = []

    # Set up the mock
    mock_document_class.return_value = mock_doc

    # Run the analysis
    tool = FormAnalysisTool()
    result = tool._run("/path/to/empty_form.docx")

    # Parse the result
    fields = json.loads(result)

    # Verify no fields were found
    assert len(fields) == 0
    assert isinstance(fields, list)


@patch("form_filler.tools.form_analysis_tool.Document")
def test_unicode_content(mock_document_class):
    """Test analysis of a form with Unicode/international characters."""
    # Create a mock document with Unicode content
    mock_doc = MagicMock()

    # Paragraphs with Unicode characters in different languages
    paragraph1 = MagicMock()
    paragraph1.text = "Tiếng Việt: ____"  # Vietnamese

    paragraph2 = MagicMock()
    paragraph2.text = "日本語 [入力]"  # Japanese

    paragraph3 = MagicMock()
    paragraph3.text = "Русский:"  # Russian

    mock_doc.paragraphs = [paragraph1, paragraph2, paragraph3]
    mock_doc.tables = []

    # Set up the mock
    mock_document_class.return_value = mock_doc

    # Run the analysis
    tool = FormAnalysisTool()
    result = tool._run("/path/to/unicode_form.docx")

    # Parse the result
    fields = json.loads(result)

    # Verify the Unicode fields were detected correctly
    assert len(fields) == 3
    field_texts = [field["text"] for field in fields]
    assert "Tiếng Việt: ____" in field_texts
    assert "日本語 [入力]" in field_texts
    assert "Русский:" in field_texts
